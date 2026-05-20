"""
Generates: TDX2026-MCP-A2A-Implementation-Evidence.docx
Author   : Md Helal — TDX 2026 Session 1557
Run with : python3.11 build_evidence_doc.py
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCREENSHOTS = os.path.join(os.path.dirname(__file__), "..", "screenshots")
OUT_FILE    = os.path.join(os.path.dirname(__file__), "TDX2026-MCP-A2A-Implementation-Evidence.docx")

# colour palette
BLUE_DARK   = RGBColor(0x00, 0x40, 0x7A)   # MuleSoft navy
BLUE_MID    = RGBColor(0x00, 0x6A, 0xD0)   # accent
GREEN_OK    = RGBColor(0x1A, 0x7A, 0x3C)
RED_ERR     = RGBColor(0xC0, 0x20, 0x20)
GREY_BG     = RGBColor(0xF2, 0xF4, 0xF7)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x1A, 0x1A, 0x1A)


# helpers

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_run(para, text, bold=False, italic=False,
            size=None, color=None, font="Calibri"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def heading(doc, text, level=1, color=BLUE_DARK):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        run.font.name = "Calibri"
    return p

def body(doc, text, space_before=0, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p

def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(1)
    p.paragraph_format.space_before  = Pt(4)
    p.paragraph_format.space_after   = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)
    # light grey background via XML shading on the paragraph
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F4F7")
    pPr.append(shd)
    return p

def insert_image(doc, filename, width=Inches(5.8), caption=None):
    path = os.path.join(SCREENSHOTS, filename)
    if not os.path.exists(path):
        body(doc, f"[Image not found: {filename}]")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(path, width=width)
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        for run in cp.runs:
            run.italic   = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.name = "Calibri"

def metric_table(doc, rows):
    """rows = list of (label, value, value_color)"""
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    for cell in hdr:
        set_cell_bg(cell, "00407A")
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.name = "Calibri"
                run.font.size = Pt(10)
    # data rows
    for i, (label, value, vcolor) in enumerate(rows):
        cells = table.rows[i + 1].cells
        cells[0].text = label
        cells[1].text = ""
        run = cells[1].paragraphs[0].add_run(value)
        run.bold = True
        run.font.color.rgb = vcolor
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        if i % 2 == 0:
            set_cell_bg(cells[0], "F2F4F7")
            set_cell_bg(cells[1], "F2F4F7")
        for cell in [cells[0], cells[1]]:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
    doc.add_paragraph()   # spacer

def evidence_table(doc, header_row, data_rows):
    """Generic evidence summary table."""
    cols   = len(header_row)
    table  = doc.add_table(rows=len(data_rows) + 1, cols=cols)
    table.style = "Table Grid"
    hcells = table.rows[0].cells
    for j, h in enumerate(header_row):
        hcells[j].text = h
        set_cell_bg(hcells[j], "00407A")
        for para in hcells[j].paragraphs:
            for run in para.runs:
                run.bold           = True
                run.font.color.rgb = WHITE
                run.font.name      = "Calibri"
                run.font.size      = Pt(10)
    for i, row in enumerate(data_rows):
        cells = table.rows[i + 1].cells
        for j, val in enumerate(row):
            cells[j].text = val
            if i % 2 == 0:
                set_cell_bg(cells[j], "F2F4F7")
            for para in cells[j].paragraphs:
                for run in para.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
    doc.add_paragraph()


# ════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ════════════════════════════════════════════════════════════

doc = Document()

# page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)


# ════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════

doc.add_paragraph()   # top padding

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(title_p,
        "Customer Complaint Resolution Network\n",
        bold=True, size=22, color=BLUE_DARK)
add_run(title_p,
        "MCP + A2A with MuleSoft Agent Fabric",
        bold=True, size=16, color=BLUE_MID)

doc.add_paragraph()

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(sub_p,
        "Implementation Evidence & Technical Documentation\n",
        bold=True, size=13, color=BLACK)
add_run(sub_p,
        "Salesforce TDX 2026 — Session 1557\n"
        "San Francisco, CA  |  April 16, 2026",
        size=11, color=RGBColor(0x44, 0x44, 0x44))

doc.add_paragraph()

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(meta_p,
        "Author: Md Helal\n"
        "Org: Student13  |  Environment: Sandbox\n"
        "All 5 Exercises Completed and Verified",
        size=11, italic=True, color=RGBColor(0x33, 0x33, 0x33))

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ════════════════════════════════════════════════════════════

heading(doc, "1. Overview", level=1)

body(doc,
     "This document provides photographic and data-driven evidence for all five hands-on "
     "exercises completed during TDX 2026 Session 1557: Connect & Govern MCP and A2A with "
     "MuleSoft Agent Fabric. Each section presents the relevant screenshots, extracted "
     "metrics, and technical interpretation of what the evidence demonstrates. The document "
     "is structured to serve both as a portfolio showcase and as a reference appendix for "
     "academic publication.")

heading(doc, "Exercise Summary", level=2)

evidence_table(doc,
    ["Exercise", "Objective", "Key Result", "Status"],
    [
        ["1", "Add A2A PII Detector policy asset to project",     "Policy asset selected from Exchange",                  "Complete"],
        ["2", "Publish agent network to Anypoint Exchange",        "36 assets published via Anypoint CLI",                 "Complete"],
        ["3", "End-to-end refund test via Slack",                  "200 OK · REF000002 · Expedited refund",                "Complete"],
        ["4", "Apply & verify A2A PII Detector policy",           "401 Unauthorized on SSN input",                        "Complete"],
        ["5a", "Agent Visualizer — live topology",                 "All 4 nodes visible · Active status · Flex Gateway",   "Complete"],
        ["5b", "Distributed trace — success path",                 "200 OK · 19 spans · 20.81s total",                     "Complete"],
        ["5c", "Distributed trace — PII rejection",                "401 · 1 span · 1.01s · No LLM call",                   "Complete"],
        ["5d", "Log search from trace",                            "102 log hits · Structured JSON logs",                  "Complete"],
    ]
)

heading(doc, "Architecture Summary", level=2)

body(doc,
     "The deployed network consists of one Broker (LLM orchestrator), one A2A sub-agent, "
     "and two MCP servers, all routed through Flex Gateway on CloudHub 2.0:")

code_block(doc,
    "Client (Slack / direct A2A POST)\n"
    "   │  JSON-RPC 2.0  ·  method: message/send\n"
    "   ▼\n"
    "Flex Gateway  —  A2A PII Detector v1.0.1  (US SSN → 401)\n"
    "   │\n"
    "   ▼\n"
    "complaint-resolution-broker  (Google Gemini 2.5 Flash Lite)\n"
    "   ├── A2A ──► Customer Verification Agent  (Node.js · Heroku /a2a/)\n"
    "   ├── MCP ──► Inventory MCP Server         (streamableHttp · Heroku /inventory)\n"
    "   │              create_return_request  ·  generate_return_label  ·  check_inventory_status\n"
    "   └── MCP ──► Finance MCP Server           (streamableHttp · Heroku /finance)\n"
    "                  process_refund  ·  create_ledger_entry")

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 2. EXERCISE 1 — Add PII Policy Asset
# ════════════════════════════════════════════════════════════

heading(doc, "2. Exercise 1 — Add A2A PII Detector Policy Asset", level=1)

body(doc,
     "The first step was to add the A2A PII Detector policy as a dependency in the agent "
     "network project. This is done inside Anypoint Code Builder by searching Anypoint "
     "Exchange and selecting the policy asset, which appends a dependency entry to "
     "exchange.json and makes the policy available for reference in agent-network.yaml.")

insert_image(doc,
    "Exercise 4 – Task 2-1.jpg",
    width=Inches(5.2),
    caption="Figure 1 — Anypoint Code Builder: 'Add Exchange Assets' dialog. "
            "A2A PII Detector v1.0.1 (MuleSoft Organization) selected.")

heading(doc, "Evidence Details", level=2)
body(doc, "The screenshot confirms:")

for item in [
    "Three policy assets are available: A2A PII Detector, MCP PII Detector, A2A Schema Validation.",
    "A2A PII Detector v1.0.1 (released Oct 3, 2025 by MuleSoft Organization) is checked.",
    "This selection adds the dependency block shown below to exchange.json.",
]:
    p = doc.add_paragraph(style="List Bullet")
    add_run(p, item, size=11)

doc.add_paragraph()
code_block(doc,
    '"dependencies": [\n'
    '  {\n'
    '    "classifier": "schema",\n'
    '    "packaging": "zip",\n'
    '    "groupId": "68ef9520-24e9-4cf2-b2f5-620025690913",\n'
    '    "assetId": "a-two-a-pii-detector",\n'
    '    "version": "1.0.1"\n'
    '  }\n'
    ']')

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 3. EXERCISE 2 — Publish to Anypoint Exchange
# ════════════════════════════════════════════════════════════

heading(doc, "3. Exercise 2 — Publish to Anypoint Exchange", level=1)

body(doc,
     "The agent network was published to Anypoint Exchange using the Anypoint CLI. "
     "A single command reads agent-network.yaml and exchange.json and publishes all "
     "constituent assets: the broker, sub-agents, MCP servers, LLM provider, the Mule "
     "application, and the top-level agent-network descriptor.")

insert_image(doc,
    "Exercise 4 – Task 2-5.jpg",
    width=Inches(5.8),
    caption="Figure 2 — Anypoint Code Builder terminal: 36 assets published to Exchange.")

heading(doc, "Key Output Lines", level=2)

code_block(doc,
    "Publishing 'agent'        'Customer Verification Agent'   → verification-agent  v1.1\n"
    "Publishing 'mcp'          'Finance MCP Server'            → finance-mcp-server  v1.1\n"
    "Publishing 'mcp'          'Inventory MCP Server'          → inventory-mcp-server v1.1\n"
    "Publishing 'llm'          'Google Gemini'                 → google-gemini v1.1\n"
    "Publishing 'app'          'customer-complaint-resolution-mule-application'\n"
    "Publishing 'agent-network''Customer Complaint Resolution Network'\n"
    "36 Assets from agent network 'customer-complaint-resolution' were published.\n"
    "    Group ID : e3ef37b3-16d8-493a-a57b-35b2d6ff2c9b\n"
    "    Published: https://anypoint.mulesoft.com/exchange/")

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 4. EXERCISE 3 — End-to-End Refund via Slack
# ════════════════════════════════════════════════════════════

heading(doc, "4. Exercise 3 — End-to-End Refund Test via Slack", level=1)

body(doc,
     "With the network deployed, the end-to-end flow was tested using RefundBot in Slack. "
     "A natural-language complaint triggers the full 5-step orchestration: customer "
     "verification via A2A, return request creation, shipping label generation, refund "
     "processing, and ledger entry — all within a single agent network invocation.")

insert_image(doc,
    "Chat-success-with-agents.jpg",
    width=Inches(5.8),
    caption="Figure 3 — Slack: RefundBot returns Refund ID REF000002 and Shipping Label URL "
            "in response to a natural-language return request.")

insert_image(doc,
    "Chat-success-with-agents-more.jpg",
    width=Inches(5.8),
    caption="Figure 4 — Slack thread: full multi-turn conversation showing the complete "
            "complaint resolution flow with expedited refund confirmation.")

heading(doc, "Verified Response Content", level=2)

code_block(doc,
    'User:    "I need to return the Laptop Stand for customer CUST002"\n\n'
    'Bot:     "Refund processed for CUST002 for order ORD003.\n'
    '          The refund ID is REF000002 and the return shipping label\n'
    '          can be found at https://shipping.example.com/labels/RET000002.pdf.\n'
    '          The refund is expedited and expected to be completed by 2026-04-21."')

heading(doc, "5-Step Orchestration Confirmed", level=2)

evidence_table(doc,
    ["Step", "Tool / Agent", "Result"],
    [
        ["1", "verification-agent  (A2A)",           "Customer CUST002 verified · expedited: true (loyalty tier)"],
        ["2", "create_return_request  (MCP Inventory)", "Return created · returnId: RET000002"],
        ["3", "generate_return_label  (MCP Inventory)", "Label URL: .../RET000002.pdf"],
        ["4", "process_refund  (MCP Finance)",        "Refund approved · refundId: REF000002 · expedited"],
        ["5", "create_ledger_entry  (MCP Finance)",   "Audit record written · linked to transaction"],
    ]
)

body(doc,
     "Note: Steps 3 and 4 run in parallel. The LLM orchestrator recognised the "
     "independence of the return label and refund operations and executed them "
     "concurrently, reducing total response time on the critical path. This was "
     "confirmed in the distributed trace (Section 6).")

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 5. EXERCISE 4 — A2A PII Detector Policy
# ════════════════════════════════════════════════════════════

heading(doc, "5. Exercise 4 — A2A PII Detector Policy", level=1)

body(doc,
     "The A2A PII Detector policy was applied to the broker via the Flex Gateway, "
     "configured to detect US Social Security Numbers and reject the message with "
     "HTTP 401 before it reaches the LLM. This section documents the policy "
     "configuration and the verified rejection behaviour.")

heading(doc, "5.1  Policy Configuration in agent-network.yaml", level=2)

body(doc, "The following YAML block was added to the broker spec:")

code_block(doc,
    "policies:\n"
    "  - ref:\n"
    "      name: a-two-a-pii-detector\n"
    "      namespace: 68ef9520-24e9-4cf2-b2f5-620025690913\n"
    "    configuration:\n"
    "      entities:\n"
    "        - US SSN\n"
    "      action: Reject    # blocks message; returns 401 to caller")

for fname, cap in [
    ("Exercise 4 – Task 2-2.jpg", "Figure 5 — YAML policy snippet added to agent-network.yaml in Anypoint Code Builder."),
    ("Exercise 4 – Task 2-3.jpg", "Figure 6 — exchange.json dependency block for a-two-a-pii-detector added."),
    ("Exercise 4 – Task 2-6.jpg", "Figure 7 — Anypoint API Manager: A2A PII Detector policy active on the broker endpoint."),
    ("Exercise 4 – Task 2-7.jpg", "Figure 8 — Policy detail: entity = US SSN, action = Reject."),
]:
    insert_image(doc, fname, width=Inches(5.5), caption=cap)

heading(doc, "5.2  PII Rejection Verified", level=2)

body(doc,
     "A test message containing the SSN '000-11-1111' was submitted. "
     "The policy fired at the Flex Gateway ingress layer before the message "
     "reached the Mule flow or the Gemini LLM:")

for fname, cap in [
    ("Exercise 4 – Task 2-8.jpg", "Figure 9 — Test request containing SSN 000-11-1111 submitted via Postman / Slack."),
    ("Exercise 4 – Task 2-9.jpg", "Figure 10 — HTTP 401 Unauthorized response. PII blocked at gateway."),
]:
    insert_image(doc, fname, width=Inches(5.5), caption=cap)

heading(doc, "5.3  LLM Proxy Policies (Google Gemini)", level=2)

body(doc,
     "The Google Gemini LLM Proxy in Anypoint API Manager was configured with "
     "two inbound policies for observability:")

for fname, cap in [
    ("google-gmi-polices-add-1.jpg", "Figure 11 — API Manager: LLM Proxies → Google Gemini → Policies overview."),
    ("google-gmi-polices-add-3.jpg", "Figure 12 — Inbound policies: Tracing (TROUBLESHOOTING) and Agent Connection Telemetry (CUSTOM)."),
]:
    insert_image(doc, fname, width=Inches(5.5), caption=cap)

body(doc,
     "Important: The yellow warning banner in API Manager reads: This instance comes "
     "from an agent network project. Policy changes made here will be reverted next time "
     "it gets redeployed. This confirms that agent-network.yaml is the authoritative "
     "source of truth for all policy configuration.")

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 6. EXERCISE 5 — Agent Visualizer & Distributed Tracing
# ════════════════════════════════════════════════════════════

heading(doc, "6. Exercise 5 — Agent Visualizer & Distributed Tracing", level=1)

body(doc,
     "Exercise 5 verified observability across two dimensions: the Agent Visualizer "
     "provided a real-time topological view of the deployed network, while Anypoint "
     "Monitoring Traces provided a span-level audit trail of individual requests.")


# 6.1 Agent Visualizer

heading(doc, "6.1  Agent Visualizer — Live Network Topology", level=2)

body(doc,
     "The Agent Visualizer at anypoint.mulesoft.com/visualizer/agentvisualizer/ displays "
     "the live topology of the deployed agent network. Filters were set to Business Group "
     "'Student13', Environment 'Sandbox', Activity Period 'Last 7 days'.")

insert_image(doc,
    "Ex5-Student-Task1-1-AgentVisualizer-Initial.jpg",
    width=Inches(5.8),
    caption="Figure 13 — Agent Visualizer: full topology graph. Complaint Resolution Broker "
            "(Agent/MuleSoft) connected to Customer Verification Agent (Agent/Customnodejs), "
            "Inventory MCP Server (MCP), and Finance MCP Server (MCP).")

insert_image(doc,
    "Exercise-5-Task-1-2-Agent-Visualizer-Broker-Details.jpg",
    width=Inches(5.8),
    caption="Figure 14 — Agent Visualizer: Broker detail panel confirming Active status, "
            "policies (A2A Agent Card, A2A PII Detector), and Flex Gateway governance.")

heading(doc, "Broker Detail Panel — Extracted Data", level=3)

evidence_table(doc,
    ["Field", "Value"],
    [
        ["Status",                 "Active"],
        ["Business Group",         "Student13"],
        ["Environment (Type)",     "Sandbox (Sandbox)"],
        ["Based on Asset Version", "v1.1.1"],
        ["A2A Policies",           "A2A Agent Card  ·  A2A PII Detector"],
        ["Transformation",         "Header Injection"],
        ["Troubleshooting",        "Tracing"],
        ["Governance & Security",  "Managed and secured with Flex Gateway"],
    ]
)


# 6.2 Traces Overview

heading(doc, "6.2  Traces Overview — All Spans", level=2)

insert_image(doc,
    "Exercise-5-Task-2-0-All-Traces-List.jpg",
    width=Inches(5.8),
    caption="Figure 15 — Anypoint Monitoring Traces Overview: full span list for the "
            "Sandbox environment. All rows show 200 OK status with timestamps from "
            "22 Apr 2026 at 12:42 am.")

insert_image(doc,
    "Exercise-5-Task-2-1-Traces-List-Broker-Filtered.jpg",
    width=Inches(5.8),
    caption="Figure 16 — Traces filtered by entity 'complaint-resolution-broker', "
            "showing only broker-rooted traces.")


# 6.3 Success Trace

heading(doc, "6.3  Distributed Trace — Successful 5-Step Run", level=2)

body(doc,
     "The trace for a successful complaint resolution request confirms that all five "
     "orchestration steps executed, with LLM reasoning spans interleaved between "
     "agent and tool calls.")

insert_image(doc,
    "Exercise-5-Task-2-2-Trace-Success-Full-Flow.jpg",
    width=Inches(5.8),
    caption="Figure 17 — Trace Detail: Trace ID ddecffb6c76a554cee11cc53da7b35bc. "
            "200 OK · 20.81s · 19 spans. Full waterfall showing Agent, LLM, and MCP Server spans.")

heading(doc, "Verified Trace Metrics", level=3)

metric_table(doc, [
    ("Trace ID",            "ddecffb6c76a554cee11cc53da7b35bc",  BLACK),
    ("Status Code",         "200 OK",                            GREEN_OK),
    ("Total Response Time", "20.81 seconds",                     BLUE_MID),
    ("Trace Start",         "22 Apr 2026 at 12:42 am",           BLACK),
    ("Root Span",           "[Agent] complaint-resolution-broker", BLACK),
    ("Total Spans",         "19",                                 BLUE_MID),
])

heading(doc, "Span Hierarchy (key spans)", level=3)

code_block(doc,
    "[Agent]      complaint-resolution-broker          20.81s\n"
    "  mule:flow                                       12.60s\n"
    "    [BROKER]  Complaint_Resolution_Broker         12.50s\n"
    "      [Agent] verification-agent                  255.79ms  ← Step 1: verify\n"
    "      [LLM]   google-gemini                       505.27ms  ← LLM reasoning\n"
    "      [Agent] verification-agent                  1.69s     ← A2A response\n"
    "      [LLM]   google-gemini                       539.27ms  ← LLM plans steps 2-5\n"
    "      [MCP Server] inventory-mcp-server           77.38ms   ← Step 2: create return\n"
    "      [LLM]   google-gemini                       833.73ms  ← Steps 3+4 (parallel)\n"
    "      [MCP Server] inventory-mcp-server           75.77ms   ← Step 3: return label\n"
    "      [LLM]   google-gemini                       551.71ms  ← Step 5: ledger\n")

body(doc,
     "Observation: Steps 3 (generate_return_label) and 4 (process_refund) produced "
     "overlapping spans in the trace waterfall, confirming that the Gemini orchestrator "
     "executed them in parallel as instructed. The Finance MCP Server spans appear "
     "within the same time window as the final Inventory MCP Server call.")


# 6.4 PII Rejection Trace

heading(doc, "6.4  Distributed Trace — PII Rejection (401)", level=2)

body(doc,
     "A message containing a US SSN ('000-11-1111') was submitted to verify the "
     "PII Detector policy. The trace demonstrates that the message was stopped at "
     "the Flex Gateway ingress layer — the Gemini LLM was never invoked.")

insert_image(doc,
    "Exercise-5-Task-2-3-Trace-Error-PII-401.jpg",
    width=Inches(5.8),
    caption="Figure 18 — Trace Detail: 401 Unauthorized. 1.01s total. Only 1 span — "
            "no LLM or MCP Server spans. SSN blocked at gateway before reaching the LLM.")

heading(doc, "PII Rejection Metrics vs. Success Metrics", level=3)

evidence_table(doc,
    ["Metric", "Success (clean message)", "PII Rejection (SSN in message)"],
    [
        ["HTTP Status",         "200 OK",           "401 Unauthorized"],
        ["Total Response Time", "20.81s",            "1.01s"],
        ["Span Count",          "19",                "1"],
        ["LLM Spans",           "4  (google-gemini)", "0  — LLM never called"],
        ["MCP Server Spans",    "2  (inventory + finance)", "0  — No tool calls"],
        ["Trace Start",         "22 Apr 2026 12:42 am", "20 Apr 2026 04:00 pm"],
    ]
)

body(doc,
     "This comparison is central to the governance argument: centralised enforcement "
     "at the gateway layer is categorically different from filtering inside agent code. "
     "The SSN cannot be memorised, logged, or accidentally forwarded by the LLM because "
     "the request is terminated at the network boundary before the Mule flow begins.")


# 6.5 Log Search

heading(doc, "6.5  Log Search — Structured Logs from Trace", level=2)

body(doc,
     "From the trace detail view, structured application logs were retrieved via "
     "Anypoint Monitoring Log Search, confirming that all runtime events are "
     "searchable, filterable, and correlated to their parent trace.")

insert_image(doc,
    "Exercise-5-Task-2-4-Logs-From-Trace.jpg",
    width=Inches(5.8),
    caption="Figure 19 — Anypoint Log Search: 102 hits linked to the trace. "
            "Structured JSON log entries from the Mule worker, class INSECURE-LOGGING.")

heading(doc, "Extracted Log Metadata", level=3)

evidence_table(doc,
    ["Field", "Value"],
    [
        ["Application",    "customer-complaint-resolution"],
        ["Log Count",      "102 hits"],
        ["Logger",         "INSECURE-LOGGING"],
        ["Log Level",      "DEBUG"],
        ["Worker ID",      "customer-complaint-resolution-00b4bc746-hl15u"],
        ["Timestamp",      "22 Apr 2026, 00:42"],
        ["Log Class",      "MuleRuntimeJwt...[Customer-complaint-resolution.uber...]"],
    ]
)

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 7. TECHNICAL NOTES FOR PUBLICATION
# ════════════════════════════════════════════════════════════

heading(doc, "7. Technical Notes for Scientific Publication", level=1)

body(doc,
     "This section summarises the novel technical observations from the implementation "
     "that are relevant for academic reference.")

points = [
    ("MCP Transport",
     "MuleSoft Agent Fabric uses streamableHttp (not HTTP+SSE) as the MCP transport "
     "for tool server connections. This is a bidirectional streaming HTTP protocol "
     "that supports server-side push without requiring SSE infrastructure."),

    ("Schema Flattening Workaround",
     "A known limitation in MuleSoft's ToolUtils.java prevents nested object properties "
     "in MCP tool input schemas. Customer address fields were therefore declared as six "
     "separate top-level strings (customerAddressName, customerAddressStreet, etc.) "
     "rather than a nested address object. This represents an active platform constraint "
     "as of April 2026."),

    ("OpenAI Strict Mode Compatibility",
     "All tool schemas use additionalProperties: false and declare every property in "
     "required[], including a contextId field. This satisfies OpenAI Strict Mode "
     "validation used by the Gemini gateway adapter. The contextId also serves as a "
     "distributed trace correlation key, linking all spans in a single orchestration run."),

    ("Gateway-Level PII Enforcement",
     "The A2A PII Detector policy operates at the Flex Gateway ingress layer, prior to "
     "the Mule flow execution. This is evidenced by the trace comparison: a rejected "
     "message produces a single span (no mule:flow child span), while a successful "
     "message produces 19 spans including flow execution, LLM calls, and MCP tool calls. "
     "The LLM cannot observe PII because the request never enters the runtime."),

    ("Parallel Orchestration",
     "The Gemini LLM orchestrator, when instructed that two steps are independent, "
     "executes them concurrently. This was confirmed by overlapping MCP Server and LLM "
     "spans in the trace waterfall for steps 3 (generate_return_label) and 4 "
     "(process_refund). This demonstrates that LLM-based orchestration can reason "
     "about causal dependencies and optimise execution paths without explicit parallel "
     "execution constructs."),

    ("Agent Visualizer Governance Confirmation",
     "The Anypoint Agent Visualizer displays active governance policies on each broker "
     "node. The presence of the 'A2A PII Detector' badge and 'Managed and secured with "
     "Flex Gateway' label in the detail panel confirms that governance is structurally "
     "enforced — not just configured but verifiably active in the deployed runtime."),
]

for title, desc in points:
    p = doc.add_paragraph(style="List Bullet")
    add_run(p, title + ": ", bold=True, size=11, color=BLUE_DARK)
    add_run(p, desc, size=11)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()


# ════════════════════════════════════════════════════════════
# 8. FIGURE INDEX
# ════════════════════════════════════════════════════════════

heading(doc, "8. Figure Index", level=1)

evidence_table(doc,
    ["Figure", "Description", "Source Screenshot"],
    [
        ["1",  "Add Exchange Assets dialog — A2A PII Detector selected",              "Exercise 4 – Task 2-1.jpg"],
        ["2",  "Terminal: 36 assets published to Exchange",                            "Exercise 4 – Task 2-5.jpg"],
        ["3",  "Slack: RefundBot returns REF000002 + Shipping Label URL",              "Chat-success-with-agents.jpg"],
        ["4",  "Slack thread: full resolution conversation",                           "Chat-success-with-agents-more.jpg"],
        ["5",  "YAML policy snippet in Anypoint Code Builder",                         "Exercise 4 – Task 2-2.jpg"],
        ["6",  "exchange.json dependency block added",                                 "Exercise 4 – Task 2-3.jpg"],
        ["7",  "API Manager: PII Detector policy active",                              "Exercise 4 – Task 2-6.jpg"],
        ["8",  "Policy detail: entity=US SSN, action=Reject",                         "Exercise 4 – Task 2-7.jpg"],
        ["9",  "Test request containing SSN 000-11-1111",                              "Exercise 4 – Task 2-8.jpg"],
        ["10", "HTTP 401 Unauthorized response",                                       "Exercise 4 – Task 2-9.jpg"],
        ["11", "API Manager: LLM Proxies → Google Gemini → Policies",                 "google-gmi-polices-add-1.jpg"],
        ["12", "Inbound policies: Tracing + Agent Connection Telemetry",               "google-gmi-polices-add-3.jpg"],
        ["13", "Agent Visualizer: full topology graph",                                "Ex5-Student-Task1-1-AgentVisualizer-Initial.jpg"],
        ["14", "Agent Visualizer: Broker detail panel",                                "Exercise-5-Task-1-2-Agent-Visualizer-Broker-Details.jpg"],
        ["15", "Traces Overview: all spans list",                                      "Exercise-5-Task-2-0-All-Traces-List.jpg"],
        ["16", "Traces filtered by complaint-resolution-broker",                       "Exercise-5-Task-2-1-Traces-List-Broker-Filtered.jpg"],
        ["17", "Trace Detail: 200 OK · 20.81s · 19 spans",                            "Exercise-5-Task-2-2-Trace-Success-Full-Flow.jpg"],
        ["18", "Trace Detail: 401 Unauthorized · 1.01s · 1 span",                     "Exercise-5-Task-2-3-Trace-Error-PII-401.jpg"],
        ["19", "Log Search: 102 hits from trace",                                      "Exercise-5-Task-2-4-Logs-From-Trace.jpg"],
    ]
)


# ════════════════════════════════════════════════════════════
# FOOTER — document info
# ════════════════════════════════════════════════════════════

doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(footer_p,
        "TDX 2026 · Session 1557 · Md Helal · Student13 / Sandbox · April 2026",
        size=9, italic=True, color=RGBColor(0x88, 0x88, 0x88))


# ════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════

doc.save(OUT_FILE)
print(f"Document saved: {OUT_FILE}")
